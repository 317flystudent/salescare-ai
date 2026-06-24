from pathlib import Path
from zipfile import ZipFile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "reports" / "outputs"
ASSET_DIR = ROOT / "reports" / "generated_assets"
DOCX_PATH = OUTPUT_DIR / "悦行销售售后AI聊天机器人开发项目报告.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(32, 37, 31)
MUTED = RGBColor(102, 112, 100)
GREEN = RGBColor(31, 122, 90)
WARM = RGBColor(242, 112, 89)
GOLD = RGBColor(255, 209, 102)
TABLE_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E0D7"


def font_path():
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def pil_font(size, bold=False):
    path = font_path()
    if path:
        return ImageFont.truetype(path, size=size, index=0)
    return ImageFont.load_default()


def draw_centered(draw, box, text, font, fill="#20251F", spacing=6):
    x1, y1, x2, y2 = box
    lines = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            trial = current + ch
            if draw.textbbox((0, 0), trial, font=font)[2] <= (x2 - x1 - 28):
                current = trial
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(heights) + spacing * max(len(lines) - 1, 0)
    y = y1 + ((y2 - y1 - total_h) / 2)
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += h + spacing


def rounded_box(draw, box, fill, outline, radius=16, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#667064", width=3):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        points = [(x2, y2), (x2 - 12, y2 - 7), (x2 - 12, y2 + 7)]
    else:
        points = [(x2, y2), (x2 + 12, y2 - 7), (x2 + 12, y2 + 7)]
    draw.polygon(points, fill=fill)


def generate_architecture_diagram(path):
    img = Image.new("RGB", (1500, 880), "#F7F9F6")
    draw = ImageDraw.Draw(img)
    title_font = pil_font(34)
    box_font = pil_font(24)
    small_font = pil_font(18)

    draw.text((60, 38), "系统技术架构图", font=title_font, fill="#20251F")
    draw.text((60, 84), "前端、后端、数据库、知识库检索与大模型服务协同工作", font=small_font, fill="#667064")

    boxes = {
        "user": (70, 190, 330, 330),
        "frontend": (430, 170, 720, 350),
        "backend": (820, 170, 1130, 350),
        "db": (1185, 470, 1435, 610),
        "retrieval": (820, 470, 1130, 610),
        "llm": (430, 520, 720, 700),
        "admin": (70, 520, 330, 700),
    }
    rounded_box(draw, boxes["user"], "#FFFFFF", "#BFCBB9")
    rounded_box(draw, boxes["frontend"], "#E3F1EA", "#1F7A5A")
    rounded_box(draw, boxes["backend"], "#FFFFFF", "#1F7A5A")
    rounded_box(draw, boxes["db"], "#FFF7D8", "#B08A00")
    rounded_box(draw, boxes["retrieval"], "#FFFFFF", "#F27059")
    rounded_box(draw, boxes["llm"], "#FFF1EE", "#F27059")
    rounded_box(draw, boxes["admin"], "#FFFFFF", "#BFCBB9")

    labels = {
        "user": "客户/客服人员\n浏览器访问",
        "frontend": "Web 前端\nHTML/CSS/JavaScript\n聊天与知识库管理",
        "backend": "后端 API\nPython HTTP Server\n会话、CRUD、AI 代理",
        "db": "SQLite / MySQL 数据库\nknowledge_base/demo_orders\nhandoff_tickets/sessions/messages",
        "retrieval": "检索与意图识别\n关键词 + 相似度评分\nTop-K 知识片段",
        "llm": "DeepSeek API\n提示词工程\n专业客服回复",
        "admin": "管理员操作\n维护知识库\n查看会话历史",
    }
    for key, label in labels.items():
        draw_centered(draw, boxes[key], label, box_font)

    arrow(draw, (330, 260), (430, 260))
    arrow(draw, (720, 260), (820, 260))
    arrow(draw, (975, 350), (975, 470))
    arrow(draw, (1130, 540), (1185, 540))
    arrow(draw, (820, 540), (720, 610))
    arrow(draw, (330, 610), (430, 610))
    arrow(draw, (975, 470), (975, 350))
    arrow(draw, (1130, 260), (1310, 470))

    draw.text((370, 236), "REST/JSON", font=small_font, fill="#667064")
    draw.text((745, 236), "CORS API", font=small_font, fill="#667064")
    draw.text((1010, 407), "构造上下文", font=small_font, fill="#667064")
    draw.text((1210, 392), "持久化", font=small_font, fill="#667064")
    draw.text((606, 486), "可选真实 LLM", font=small_font, fill="#667064")

    img.save(path)


def generate_er_diagram(path):
    img = Image.new("RGB", (1500, 940), "#F7F9F6")
    draw = ImageDraw.Draw(img)
    title_font = pil_font(34)
    box_font = pil_font(22)
    mono_font = pil_font(18)

    draw.text((60, 38), "数据库 ER 图", font=title_font, fill="#20251F")
    draw.text((60, 84), "五张核心表：知识库、演示订单、人工工单、会话、消息；课堂演示可切换 MySQL", font=mono_font, fill="#667064")

    entities = {
        "sessions": (80, 180, 420, 480),
        "handoff": (80, 520, 420, 850),
        "messages": (500, 180, 840, 620),
        "knowledge": (920, 180, 1420, 430),
        "orders": (920, 485, 1420, 735),
    }
    for key, box in entities.items():
        rounded_box(draw, box, "#FFFFFF", "#BFCBB9")

    headers = {
        "sessions": ("sessions", "#E3F1EA", "#1F7A5A"),
        "handoff": ("handoff_tickets", "#FFF7D8", "#B08A00"),
        "messages": ("messages", "#E3F1EA", "#1F7A5A"),
        "knowledge": ("knowledge_base", "#FFF7D8", "#B08A00"),
        "orders": ("demo_orders", "#FFF1EE", "#F27059"),
    }
    for key, (label, fill, outline) in headers.items():
        x1, y1, x2, _ = entities[key]
        draw.rounded_rectangle((x1, y1, x2, y1 + 58), radius=16, fill=fill, outline=outline, width=2)
        draw.text((x1 + 18, y1 + 15), label, font=box_font, fill="#20251F")

    fields = {
        "sessions": [
            "id TEXT PK",
            "channel TEXT",
            "created_at TEXT",
            "updated_at TEXT",
        ],
        "handoff": [
            "ticket_no TEXT PK",
            "session_id TEXT",
            "reason TEXT",
            "priority TEXT",
            "status TEXT",
            "assigned_team TEXT",
        ],
        "messages": [
            "id INTEGER PK",
            "session_id TEXT FK",
            "role TEXT",
            "content TEXT",
            "intent TEXT",
            "metadata_json TEXT",
            "created_at TEXT",
        ],
        "knowledge": [
            "id INTEGER PK",
            "category TEXT",
            "title TEXT",
            "question TEXT",
            "answer TEXT",
            "keywords TEXT",
            "updated_at TEXT",
        ],
        "orders": [
            "order_no TEXT PK",
            "product_name TEXT",
            "status TEXT",
            "logistics_status TEXT",
            "aftersale_status TEXT",
        ],
    }
    for key, lines in fields.items():
        x1, y1, _x2, _y2 = entities[key]
        y = y1 + 88
        for line in lines:
            draw.text((x1 + 24, y), line, font=mono_font, fill="#20251F")
            y += 44

    arrow(draw, (420, 350), (500, 350), fill="#1F7A5A", width=4)
    draw.text((434, 314), "1 : N", font=mono_font, fill="#1F7A5A")
    arrow(draw, (250, 480), (250, 520), fill="#1F7A5A", width=4)
    draw.text((270, 496), "转人工", font=mono_font, fill="#1F7A5A")
    arrow(draw, (920, 305), (840, 350), fill="#F27059", width=4)
    draw.text((842, 280), "检索引用", font=mono_font, fill="#F27059")
    arrow(draw, (920, 610), (840, 500), fill="#F27059", width=4)
    draw.text((842, 585), "订单查询", font=mono_font, fill="#F27059")

    draw.text((92, 885), "说明：messages.session_id 和 handoff_tickets.session_id 关联会话；knowledge_base、demo_orders 和 handoff_tickets 共同支撑客服回复。", font=mono_font, fill="#667064")
    img.save(path)


def set_run_font(run, name="Calibri", east_asia="微软雅黑", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def paragraph_border_bottom(paragraph, color="BFCBB9", size="8"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table(doc, headers, rows, widths, caption=None):
    if caption:
        p = doc.add_paragraph(caption)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            set_run_font(run, size=10, color=MUTED, italic=True)

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    set_table_geometry(tbl, widths)
    mark_header_row(tbl.rows[0])
    for idx, header in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        shade_cell(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10.5, color=TEXT, bold=True)
    for row_data in rows:
        row = tbl.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_run_font(run, size=10, color=TEXT)
    set_table_geometry(tbl, widths)
    doc.add_paragraph()
    return tbl


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    end_run = paragraph.add_run(" 页")
    set_run_font(end_run, size=9, color=MUTED)


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("悦行销售/售后 AI 聊天机器人项目报告")
    set_run_font(left, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_para(doc, text, style=None, bold=False, color=None, size=None, align=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11, color=TEXT)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=11, color=TEXT)


def add_callout(doc, title, body):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    set_table_geometry(tbl, [9360])
    mark_header_row(tbl.rows[0])
    cell = tbl.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    title_run = p.add_run(title + "：")
    set_run_font(title_run, size=10.5, color=DARK_BLUE, bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run, size=10.5, color=TEXT)
    doc.add_paragraph()


def build_cover(doc):
    add_para(doc, "课程项目报告", size=12, color=MUTED, bold=True, after=4)
    title = add_para(
        doc,
        "悦行销售（售后）服务 AI 聊天机器人开发项目报告",
        size=24,
        color=TEXT,
        bold=True,
        after=6,
    )
    subtitle = add_para(
        doc,
        "面向电动车销售咨询、售后服务与故障排查的可交互 Web 应用",
        size=13,
        color=MUTED,
        after=16,
    )
    paragraph_border_bottom(subtitle, color="1F7A5A", size="10")

    rows = [
        ("课程名称", "大数据机器学习"),
        ("项目类型", "小组项目 / 全栈 AI 聊天机器人"),
        ("项目周期", "10 周"),
        ("小组成员/学号", "待填写"),
        ("技术栈", "原生 Web 前端、Python HTTP API、SQLite/MySQL、订单演示、人工工单、DeepSeek API 可选接入"),
        ("本地演示 URL", "http://127.0.0.1:5173"),
        ("后端 API URL", "http://127.0.0.1:8000"),
        ("生产环境 URL", "http://1.14.184.75"),
        ("开源代码仓库", "https://github.com/317flystudent/salescare-ai"),
        ("报告日期", "2026 年 6 月 10 日"),
    ]
    add_table(doc, ["项目", "内容"], rows, [1900, 7460], caption="首页交付信息")
    add_callout(
        doc,
        "说明",
        "本项目已完成可本地访问和交互的 Web 应用、后端 API、SQLite/MySQL 双模式持久化、知识库检索、演示订单查询流程与 DeepSeek API 接入代码。生产环境 URL 和开源代码仓库地址已补充；DeepSeek API Key 仅通过环境变量配置，不写入代码仓库或报告正文。",
    )
    doc.add_page_break()


def build_report():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    arch_path = ASSET_DIR / "architecture.png"
    er_path = ASSET_DIR / "er_diagram.png"
    generate_architecture_diagram(arch_path)
    generate_er_diagram(er_path)

    doc = Document()
    configure_styles(doc)
    add_header_footer(doc.sections[0])
    build_cover(doc)

    add_para(doc, "摘要", style="Heading 1")
    add_para(
        doc,
        "本项目围绕电商销售与售后服务场景，设计并实现了一个面向“悦行智能电动车”品牌的 AI 聊天机器人。系统将买家页面与商家后台分离：买家进入后看到正常的售后聊天、欢迎词和自助入口；商家可通过“商家后台”进入管理页，查看知识库、会话、演示订单、人工工单、买家订单信息和订单汇总。技术实现采用前后端分离结构：前端为响应式 Web 应用，后端提供 REST API、会话管理、知识库 CRUD、演示订单生成、SQLite/MySQL 双模式持久化与 AI 生成代理。AI 层采用“订单工作流 + 知识库检索 + 提示词约束 + 大模型可选调用”的组合方案：配置 DeepSeek API Key 后可调用真实大模型；未配置时使用本地检索式回答，保障课堂演示稳定性。",
    )
    add_callout(
        doc,
        "项目亮点",
        "系统不是简单转发用户问题，而是在回答前完成订单号识别、意图识别、客户情绪判断、Top-K 知识片段检索和边界控制；买家页提供 1-6 个电商售后入口，商家后台提供订单明细、订单汇总、随机演示订单和转人工工单，便于课堂现场展示完整业务闭环。",
    )

    add_para(doc, "1. 项目背景与市场分析", style="Heading 1")
    add_para(
        doc,
        "电动车、家电、数码产品等消费品的销售和售后服务具有高频、重复、时效要求强的特点。传统人工客服在促销咨询、物流查询、保修解释和基础故障排查上投入大量重复劳动，客户在等待过程中容易产生焦虑和不满。AI 聊天机器人适合承担标准化问题的首轮响应，并把复杂争议升级给人工客服。",
    )
    add_table(
        doc,
        ["痛点", "表现", "机器人机会"],
        [
            ("咨询重复", "同类车型、价格、促销、保修问题高频出现。", "用知识库标准答案快速响应，减轻人工客服压力。"),
            ("售后时效", "物流慢、维修慢、检测结论不透明，用户容易产生负面情绪。", "先安抚再收集订单/车型/故障信息，形成可转人工的结构化线索。"),
            ("知识分散", "门店政策、产品参数和售后条款分散在不同文档。", "统一知识库管理，支持后台增删改查。"),
            ("回答风险", "客服可能承诺超出政策范围的退款、保修或库存信息。", "通过系统提示词限制模型边界，无法确认时引导补充信息或升级人工。"),
        ],
        [1700, 3600, 4060],
        caption="市场痛点与项目机会",
    )
    add_para(doc, "1.1 目标用户画像", style="Heading 2")
    add_bullets(
        doc,
        [
            "电商消费者：关注车型差异、优惠政策、物流进度和发票开具。",
            "线下门店顾客：关注试骑、提车、门店维修预约和保修凭证。",
            "售后问题用户：关注无法启动、续航下降、App 连接失败等故障排查。",
            "客服/管理员：需要维护知识库、查看会话历史并把复杂问题转人工。",
        ],
    )
    add_para(doc, "1.2 核心功能定位", style="Heading 2")
    add_table(
        doc,
        ["功能", "面向场景", "实现方式"],
        [
            ("产品咨询", "车型对比、续航、适用人群。", "知识库检索 + 大模型组织语言。"),
            ("订单查询", "物流、自提、提车码和售后进度说明。", "识别 YX 开头订单号，从 demo_orders 表返回付款、物流和售后状态。"),
            ("退换货政策", "七天无理由、质量问题处理。", "标准政策问答与边界提示。"),
            ("故障排查", "无法启动、续航下降、App 蓝牙失败。", "步骤化排查 + 建议拍摄仪表/故障照片。"),
            ("情绪安抚", "投诉、差评、维修慢。", "情绪关键词识别，先道歉理解，再给处理路径。"),
            ("转人工工单", "质量争议、维修慢、用户要求人工介入。", "生成 RG 开头工单号，记录优先级、状态、处理小组和预计响应时间。"),
            ("知识库管理", "运营人员维护标准答案。", "前端表单 + 后端 CRUD + SQLite/MySQL。"),
        ],
        [1700, 3300, 4360],
        caption="核心功能清单",
    )

    add_para(doc, "2. 系统架构设计", style="Heading 1")
    add_para(
        doc,
        "系统采用前后端分离架构。前端负责买家聊天页、商家后台页、自助入口、用户输入、订单明细、订单汇总、知识库表单和 API 调用；后端负责业务逻辑、数据库持久化、演示订单生成、知识库检索、提示词拼接和大模型代理。该结构便于后续将前端部署到静态托管平台，将后端部署到云服务器或容器环境。",
    )
    arch_shape = doc.add_picture(str(arch_path), width=Inches(6.3))
    arch_shape._inline.docPr.set("title", "系统技术架构图")
    arch_shape._inline.docPr.set("descr", "前端、后端、数据库、知识库检索和 DeepSeek 大模型服务之间的数据流架构图。")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "图 1 系统技术架构图", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_table(
        doc,
        ["层级", "技术选型", "选择理由"],
        [
            ("前端", "HTML、CSS、JavaScript", "无需构建工具，便于课程演示；响应式布局兼容桌面和移动端。"),
            ("后端", "Python 标准库 ThreadingHTTPServer", "减少依赖安装风险，提供 REST API 与 CORS。"),
            ("数据库", "SQLite/MySQL", "默认 SQLite 保证本地稳定；课堂切换 MySQL 后可查看 knowledge_base、demo_orders、handoff_tickets、sessions、messages 五张表。"),
            ("AI 服务", "DeepSeek API 可选接入", "支持真实大模型生成；未配置 Key 时本地检索兜底。"),
            ("检索", "关键词 + 字符相似度 Top-K", "适合中文短问句和售后知识库初期规模。"),
        ],
        [1500, 2800, 5060],
        caption="技术栈选型",
    )
    add_para(doc, "2.1 核心功能模块", style="Heading 2")
    add_table(
        doc,
        ["模块", "职责", "主要文件"],
        [
            ("用户对话管理", "创建会话、保存用户与机器人消息、返回会话 ID。", "backend/server.py, backend/database.py"),
            ("演示订单管理", "生成随机演示订单，按订单号返回付款、物流和售后状态。", "backend/server.py, backend/database.py"),
            ("人工工单管理", "识别转人工请求，生成工单号，保存状态、优先级和处理小组。", "backend/server.py, backend/database.py"),
            ("知识库管理", "提供知识库列表、新增、更新、删除接口。", "backend/database.py, frontend/app.js"),
            ("AI 意图识别与响应", "识别订单号、咨询意图、客户情绪，检索知识片段并生成回复。", "backend/ai_service.py"),
            ("大模型代理", "按系统提示词和知识库上下文调用 DeepSeek。", "backend/ai_service.py"),
            ("商家后台", "与买家聊天页分离，维护知识库，查看买家订单、订单汇总、会话列表和人工工单。", "frontend/index.html, frontend/app.js"),
        ],
        [1900, 4200, 3260],
        caption="核心模块设计",
    )
    add_para(doc, "2.2 数据库 ER 设计", style="Heading 2")
    er_shape = doc.add_picture(str(er_path), width=Inches(6.3))
    er_shape._inline.docPr.set("title", "数据库 ER 图")
    er_shape._inline.docPr.set("descr", "knowledge_base、demo_orders、handoff_tickets、sessions 和 messages 五张核心表及其关系，支持 SQLite 和 MySQL。")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "图 2 数据库 ER 图", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "3. AI 模型集成与优化", style="Heading 1")
    add_para(
        doc,
        "AI 核心采用分层策略：首先由本地规则完成意图和情绪识别，再从知识库中检索 Top-K 相关条目，最后将检索内容、历史上下文和系统提示词组合后发送给 DeepSeek。该策略属于轻量 RAG（Retrieval-Augmented Generation）实现，能降低大模型幻觉风险，并提升售后政策类回答的稳定性。",
    )
    add_para(doc, "3.1 知识库构建", style="Heading 2")
    add_para(
        doc,
        "项目为虚构品牌“悦行智能电动车”构建了 12 条初始知识库，覆盖产品、价格促销、订单、退换货、保修、故障排查、维修预约、发票和投诉安抚。每条知识包含分类、标题、标准问题、标准答案和关键词，既方便人工维护，也便于检索算法打分。",
    )
    add_table(
        doc,
        ["知识类别", "示例条目", "服务价值"],
        [
            ("产品咨询", "悦行 S1、悦行 Pro 车型对比。", "辅助销售转化，减少人工重复介绍。"),
            ("售后政策", "七天无理由、整车与电池保修。", "统一口径，避免客服承诺不一致。"),
            ("故障排查", "无法启动、续航变短、App 连接失败。", "给出可执行步骤，降低门店无效到访。"),
            ("情绪安抚", "投诉、维修慢、质量争议。", "提升首轮响应体验，及时转人工。"),
        ],
        [1900, 3600, 3860],
        caption="知识库覆盖范围",
    )
    add_para(doc, "3.2 提示词工程设计", style="Heading 2")
    add_callout(
        doc,
        "系统提示词原则",
        "限定角色为“悦行智能电动车销售与售后服务 AI 助手”；要求优先依据知识库；禁止编造订单状态、价格、库存、检测结论；遇到安全风险、质量争议、多次维修或强烈不满时建议升级人工客服。",
    )
    add_numbered(
        doc,
        [
            "专业性：回答必须围绕销售/售后，不扩展到无关领域。",
            "准确性：产品参数、保修时长和退换货规则以知识库为准。",
            "边界意识：缺少订单号、城市、车型或故障证据时主动询问，不虚构结果。",
            "情绪安抚：识别“生气、投诉、差评、着急”等词后先表达理解，再给下一步。",
            "可转人工：对安全、质量争议、多次维修、强烈不满场景明确建议升级人工。",
        ],
    )
    add_para(doc, "3.3 调优过程", style="Heading 2")
    add_table(
        doc,
        ["问题", "初始风险", "优化策略"],
        [
            ("模型可能编造订单状态", "用户问物流时模型直接给虚假进度。", "提示词禁止编造，并要求用户提供订单号或手机号后四位。"),
            ("政策回答不稳定", "退换货与保修口径容易泛化。", "把政策写入知识库，并把 Top-K 片段注入提示词。"),
            ("投诉场景语气生硬", "直接给流程可能激化情绪。", "先识别负面情绪并添加安抚前缀。"),
            ("无 API Key 时无法演示", "课堂现场网络或 Key 不稳定。", "实现本地检索兜底，保证核心流程可展示。"),
        ],
        [1900, 3600, 3860],
        caption="AI 调优记录",
    )

    add_para(doc, "4. 系统实现", style="Heading 1")
    add_para(
        doc,
        "项目代码位于 E:\\dsjhomework，采用清晰目录结构组织。后端启动时会自动初始化数据库并写入种子知识库；默认使用 SQLite，课堂演示可通过环境变量切换到 MySQL。前端通过 fetch 调用 API，所有回答、意图、来源和会话 ID 都会返回给页面展示。",
    )
    add_table(
        doc,
        ["接口", "方法", "用途"],
        [
            ("/api/health", "GET", "健康检查，前端用于显示后端连接状态。"),
            ("/api/chat", "POST", "接收用户消息，返回机器人回答、意图、知识来源和会话 ID。"),
            ("/api/knowledge", "GET/POST", "查询或新增知识库。"),
            ("/api/knowledge/{id}", "PUT/DELETE", "更新或删除指定知识库条目。"),
            ("/api/demo-orders", "GET", "查看演示订单和可选演示商品。"),
            ("/api/demo-orders/create", "POST", "生成随机演示订单，用于课堂现场查询。"),
            ("/api/handoff-tickets", "GET", "查看人工客服工单。"),
            ("/api/handoff-tickets/create", "POST", "创建人工客服工单。"),
            ("/api/sessions", "GET", "查看会话列表。"),
            ("/api/sessions/{id}/messages", "GET", "查看指定会话消息历史。"),
        ],
        [2600, 1400, 5360],
        caption="后端 API 设计",
    )
    add_para(doc, "4.1 后端实现要点", style="Heading 2")
    add_bullets(
        doc,
        [
            "server.py 基于 ThreadingHTTPServer 实现 REST API，并统一处理 CORS、JSON 解析和错误响应。",
            "database.py 使用 SQLite/MySQL 保存 knowledge_base、demo_orders、handoff_tickets、sessions 和 messages，支持自动建表和种子数据初始化。",
            "ai_service.py 实现订单号识别、转人工识别、意图分类、情绪检测、知识库 Top-K 检索、DeepSeek 请求和本地兜底回答。",
            "config.py 通过环境变量管理端口、数据库模式、MySQL 连接信息、DeepSeek API Key、模型名称和超时时间。",
        ],
    )
    add_para(doc, "4.2 前端实现要点", style="Heading 2")
    add_bullets(
        doc,
        [
            "index.html 将买家聊天页和商家后台分离，买家默认只看到聊天，商家通过入口进入后台管理。",
            "styles.css 设计买家模式和后台模式，桌面端与移动端均可切换布局，避免横向滚动。",
            "app.js 管理页面模式切换、API 地址、会话 ID、本地存储、消息渲染、知识库 CRUD、演示订单生成、订单汇总、转人工工单刷新和错误提示。",
            "assets/brand-mark.svg 提供品牌视觉标识，使应用不是纯文本界面。",
        ],
    )
    add_para(doc, "4.3 关键挑战与解决方案", style="Heading 2")
    add_table(
        doc,
        ["挑战", "解决方案", "效果"],
        [
            ("不依赖复杂环境", "后端和前端均使用标准库/原生能力实现。", "降低安装失败风险，便于课堂演示。"),
            ("AI 幻觉控制", "知识库优先、提示词约束、缺信息时追问。", "售后政策类回答更可控。"),
            ("演示稳定性", "DeepSeek 可选接入，本地检索兜底。", "没有 Key 或网络时仍可完成演示。"),
            ("移动端适配", "CSS 断点、单列布局、固定按钮尺寸。", "390px 宽度测试无横向溢出。"),
        ],
        [2200, 3900, 3260],
        caption="实现挑战与处理",
    )

    add_para(doc, "5. 系统测试与部署", style="Heading 1")
    add_para(
        doc,
        "本地测试采用 API 烟测、浏览器交互测试和移动端视口检查。由于生产部署需要云服务器、域名或平台账号，本报告将生产 URL 标记为待部署后填写；项目代码已提供部署步骤和环境变量说明。",
    )
    add_table(
        doc,
        ["测试项", "输入/操作", "期望结果", "当前结果"],
        [
            ("后端健康检查", "GET /api/health", "返回 ok=true。", "通过。"),
            ("聊天接口", "电池保修多久？", "识别保修政策并引用知识库。", "通过，返回电池 36 个月保修说明。"),
            ("订单演示", "生成演示订单并发送 YX 订单号", "返回商品、付款、物流、售后状态，并写入 messages。", "通过。"),
            ("转人工工单", "点击“转人工”或输入我要转人工", "返回 RG 工单号，写入 handoff_tickets，并在右侧工单页显示。", "通过。"),
            ("故障排查", "我的车无法启动，应该怎么排查？", "给出逐步排查建议。", "通过。"),
            ("后台概览", "点击商家后台", "买家聊天页隐藏，后台显示订单总数、成交额、售后订单、工单、会话和买家订单表。", "通过。"),
            ("知识库加载", "商家后台打开知识库", "显示初始 12 条知识。", "通过。"),
            ("桌面端 UI", "1280x720 浏览器检查", "买家聊天页和商家后台模式可切换且无明显溢出。", "通过。"),
            ("移动端 UI", "390x844 视口检查", "单列布局，无横向滚动。", "通过。"),
        ],
        [1700, 2700, 3000, 1960],
        caption="功能与界面测试用例",
    )
    add_para(doc, "5.1 性能与安全测试", style="Heading 2")
    add_bullets(
        doc,
        [
            "性能：本地知识库检索规模较小，单次回答主要耗时来自大模型 API；本地兜底回答可在秒级完成。",
            "并发：后端采用 ThreadingHTTPServer，可支持多个本地浏览器请求，适合作业演示和小规模测试。",
            "安全：DeepSeek API Key 通过环境变量读取，不写入前端代码；前端只访问后端代理，避免泄露密钥。",
            "数据边界：不保存身份证、地址、支付信息等敏感数据；订单查询使用 demo_orders 演示表，人工工单使用 handoff_tickets 演示表，不对接真实第三方订单或客服系统。",
        ],
    )
    add_para(doc, "5.2 部署方案", style="Heading 2")
    add_numbered(
        doc,
        [
            "准备云服务器或平台环境，安装 Python 3.10+。",
            "上传项目代码，设置 DEEPSEEK_API_KEY、CHATBOT_HOST、CHATBOT_PORT、CHATBOT_DB 等环境变量。",
            "后端运行 backend/server.py，前端可用 Nginx、静态托管或 python -m http.server 发布。",
            "把前端 API 地址改为后端公网地址，并在页面设置中保存。",
            "完成联调后，将生产环境 URL 和代码仓库地址补入本报告首页。",
        ],
    )

    add_para(doc, "6. 项目管理与迭代记录", style="Heading 1")
    add_table(
        doc,
        ["周次", "阶段", "交付内容"],
        [
            ("第 1-2 周", "市场调研与定位", "完成销售/售后痛点、用户画像和功能范围。"),
            ("第 3 周", "系统设计", "确定前后端分离架构、数据库 ER 和模块边界。"),
            ("第 4-5 周", "AI 核心实现", "构建知识库、意图识别、提示词和 DeepSeek 接入。"),
            ("第 6-7 周", "全栈开发", "完成聊天界面、知识库管理、订单查询、人工工单、会话历史和 API。"),
            ("第 8 周", "集成测试", "完成接口、桌面端和移动端测试。"),
            ("第 9 周", "部署准备", "整理 README、环境变量和部署步骤。"),
            ("第 10 周", "报告与答辩", "形成项目报告，准备演示脚本和答辩说明。"),
        ],
        [1300, 2200, 5860],
        caption="10 周项目计划",
    )
    add_para(doc, "7. 总结与答辩准备", style="Heading 1")
    add_para(
        doc,
        "本项目完成了课程要求中的核心闭环：市场分析、系统架构、AI 核心、前后端开发、数据库持久化、测试和部署准备。系统可围绕销售/售后常见问题完成首轮客服响应，并通过知识库维护、订单工作流、转人工工单、商家后台汇总和大模型提示词工程降低回答风险。答辩演示时建议按“买家页欢迎入口 -> 订单号查询 -> 故障排查 -> 投诉安抚 -> 转人工 -> 商家后台查看订单汇总 -> MySQL 查看 messages/demo_orders/handoff_tickets”的顺序展示，能够体现功能完整性、交互体验和技术深度。",
    )
    add_callout(
        doc,
        "提交前检查",
        "报告首页已补充生产环境 URL 和代码仓库地址；如需展示 DeepSeek 真实调用，请在演示机器或云服务器环境变量中配置 DEEPSEEK_API_KEY，避免把密钥写入仓库。",
    )

    add_para(doc, "附录 A：示例对话", style="Heading 1")
    add_table(
        doc,
        ["角色", "内容"],
        [
            ("用户", "我的车无法启动，应该怎么排查？"),
            ("机器人", "请先检查电量是否充足、空气开关是否打开、刹车把是否回弹、钥匙/NFC 是否解锁成功。如果仪表亮但无法行驶，可能是刹车断电开关或控制器保护；建议拍摄仪表提示并联系门店检测。"),
            ("用户", "查询订单 YX202606090001"),
            ("机器人", "已为您查询到订单 YX202606090001：商品为悦行 S1 城市通勤版，付款状态已支付，订单状态已出库，物流进度为已到达华东转运中心，售后状态暂未发起。"),
            ("用户", "我很生气，维修太慢了。"),
            ("机器人", "非常抱歉给您带来不好的体验，我先帮您把问题拆清楚并给出下一步处理方案。若问题已经影响使用，建议同步转人工客服并保留订单、维修单或照片凭证。"),
            ("用户", "我要转人工，维修太慢了。"),
            ("机器人", "已为您转接人工客服，并创建人工工单 RG20260610xxxx。当前状态为排队中，系统会显示处理小组、优先级和预计响应时间。"),
        ],
        [1400, 7960],
        caption="演示用对话脚本",
    )

    add_para(doc, "附录 B：公有云部署截图", style="Heading 1")
    add_para(
        doc,
        "以下截图用于保留云服务器部署、数据库初始化、后端服务启动和 Nginx 对外发布的关键过程。涉及 DeepSeek Key、数据库密码、服务器密码等敏感信息的原始截图已裁剪或不展示，报告中仅保留答辩检查所需的部署证明信息。",
    )
    deployment_screenshots = [
        (
            "B.1 云服务器实例与公网地址",
            ASSET_DIR / "deployment_screenshots" / "01_cloud_server_config.png",
            "腾讯云轻量应用服务器处于运行状态，并分配了公网 IPv4 地址，可作为 Web 应用生产环境入口。",
            6.3,
        ),
        (
            "B.2 SSH 远程登录服务器",
            ASSET_DIR / "deployment_screenshots" / "02_ssh_login_success.png",
            "通过 SSH 登录 Ubuntu 服务器，进入云端部署环境。",
            6.3,
        ),
        (
            "B.3 运行环境与项目文件",
            ASSET_DIR / "deployment_screenshots" / "03_runtime_versions_and_upload.png",
            "服务器安装 Python、Nginx、MySQL 等组件，并同步项目代码、创建虚拟环境、安装 PyMySQL 依赖。",
            6.3,
        ),
        (
            "B.4 MySQL 数据库初始化验证",
            ASSET_DIR / "deployment_screenshots" / "04_mysql_tables_verified.png",
            "salescare_ai 数据库包含订单、人工工单、知识库、消息和会话表，知识库记录数为 12 条。",
            5.4,
        ),
        (
            "B.5 后端服务运行状态",
            ASSET_DIR / "deployment_screenshots" / "05_backend_service_running.png",
            "systemd 后台服务 salescare-ai 已启动，健康检查接口返回 ok，说明后端 API 可用。",
            6.3,
        ),
        (
            "B.6 Nginx 反向代理配置",
            ASSET_DIR / "deployment_screenshots" / "06_nginx_proxy_success.png",
            "Nginx 托管前端页面，并将 /api/ 请求反向代理到本机后端服务，配置测试结果为 successful。",
            6.3,
        ),
        (
            "B.7 公网访问与 AI 对话演示",
            ASSET_DIR / "deployment_screenshots" / "07_public_ai_chat_demo.png",
            "浏览器通过公网地址访问悦行销售/售后 AI 助手，并完成产品推荐类自然语言对话。当前课程演示环境使用 HTTP，正式上线可进一步配置 HTTPS 证书。",
            6.3,
        ),
        (
            "B.8 开源代码仓库上传完成",
            ASSET_DIR / "deployment_screenshots" / "08_github_repository_files.png",
            "GitHub 仓库已成功上传项目代码，包含 backend、frontend、database、docs、reports、scripts 等核心目录，并保留两次提交记录。",
            6.3,
        ),
        (
            "B.9 README 项目地址展示",
            ASSET_DIR / "deployment_screenshots" / "09_github_readme_project_urls.png",
            "仓库 README 中展示了 Web 应用生产环境 URL 与开源代码仓库地址，便于老师直接检查线上系统和源码。",
            6.3,
        ),
    ]
    for title, image_path, description, width in deployment_screenshots:
        if not image_path.exists():
            continue
        add_para(doc, title, style="Heading 2")
        add_para(doc, description)
        figure = doc.add_picture(str(image_path), width=Inches(width))
        figure._inline.docPr.set("title", title)
        figure._inline.docPr.set("descr", description)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(DOCX_PATH)
    return DOCX_PATH


def structural_audit(docx_path):
    with ZipFile(docx_path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        styles_xml = zf.read("word/styles.xml").decode("utf-8")
        rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
    checks = {
        "has_images": "image" in rels_xml,
        "heading_style_present": "Heading1" in styles_xml or "Heading 1" in styles_xml,
        "table_width_9360": 'w:w="9360"' in document_xml,
        "table_indent_120": 'w:tblInd' in document_xml and 'w:w="120"' in document_xml,
        "page_margins_1440": 'w:pgMar' in document_xml and 'w:left="1440"' in document_xml,
    }
    return checks


if __name__ == "__main__":
    path = build_report()
    print(path)
    print(structural_audit(path))
